import os
import logging
from typing import Optional
from docx import Document

logger = logging.getLogger(__name__)

def extract_text_from_file(file_path: str, filename: str) -> Optional[str]:
    """
    Extracts text content from .txt, .md, or .docx files.
    """
    ext = os.path.splitext(filename)[1].lower()
    
    try:
        if ext in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif ext in ['.docx', '.doc']:
            # Note: python-docx handles .docx. For .doc (old format), 
            # we might need another library, but for now we'll try Document().
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        else:
            logger.warning(f"Unsupported file extension: {ext}")
            return None
    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {e}")
        return None

def save_execution_plan(name: str, plan_data: dict, storage_dir: str = "storage/plans"):
    """
    Saves an execution plan as a JSON file.
    """
    if not os.path.exists(storage_dir):
        os.makedirs(storage_dir)
    
    # Ensure name ends with .json
    if not name.endswith('.json'):
        name += '.json'
        
    file_path = os.path.join(storage_dir, name)
    import json
    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump(plan_data, f, indent=4, ensure_ascii=False)

def load_execution_plan(name: str, storage_dir: str = "storage/plans") -> Optional[dict]:
    """
    Loads an execution plan from a JSON file.
    """
    if not name.endswith('.json'):
        name += '.json'
        
    file_path = os.path.join(storage_dir, name)
    if not os.path.exists(file_path):
        return None
        
    import json
    with open(file_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def list_saved_plans(storage_dir: str = "storage/plans") -> list:
    """
    Lists all saved execution plans.
    """
    if not os.path.exists(storage_dir):
        return []
        
    return [f.replace('.json', '') for f in os.listdir(storage_dir) if f.endswith('.json')]
